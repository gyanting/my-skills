#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Word报告生成脚本 v2
支持从 vuln_knowledge_base.yaml 注入深度漏洞分析
"""

import json
import sys
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, List, Optional

try:
    import yaml
    _HAS_YAML = True
except ImportError:
    _HAS_YAML = False


# ============================================================
# 漏洞知识库加载
# ============================================================
VULN_KB = None

def _load_vuln_knowledge_base() -> Dict:
    """加载漏洞知识库"""
    global VULN_KB
    if VULN_KB is not None:
        return VULN_KB

    script_dir = os.path.dirname(os.path.abspath(__file__))
    search_paths = [
        os.path.join(script_dir, '..', 'references', 'vuln_knowledge_base.yaml'),
        os.path.join(os.getcwd(), 'references', 'vuln_knowledge_base.yaml'),
        os.path.join(os.getcwd(), '.claude', 'skills', 'security-alert-analysis',
                     'references', 'vuln_knowledge_base.yaml'),
    ]
    for path in search_paths:
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    if _HAS_YAML:
                        VULN_KB = yaml.safe_load(f)
                    else:
                        VULN_KB = _parse_vuln_yaml(f.read())
                    return VULN_KB
            except Exception:
                continue
    VULN_KB = {"vulnerabilities": []}
    return VULN_KB


def _parse_vuln_yaml(content: str) -> Dict:
    """简易 YAML 解析（降级方案）"""
    vulns = []
    current = {}
    in_vuln = False
    for line in content.split('\n'):
        s = line.strip()
        if s == 'vulnerabilities:':
            continue
        if s.startswith('- name:'):
            if current:
                vulns.append(current)
            current = {'name': s[7:].strip().strip('"').strip("'")}
            in_vuln = True
        elif in_vuln and ':' in s and not s.startswith('-'):
            k, v = s.split(':', 1)
            k = k.strip()
            v = v.strip().strip('"').strip("'")
            if k == 'match_events':
                current[k] = [x.strip().strip('"').strip("'") for x in v.strip('[]').split(',')]
            else:
                current[k] = v
    if current:
        vulns.append(current)
    return {"vulnerabilities": vulns}


def _lookup_vuln(event_name: str) -> Optional[Dict]:
    """在知识库中查找漏洞信息"""
    kb = _load_vuln_knowledge_base()
    for vuln in kb.get('vulnerabilities', []):
        match_events = vuln.get('match_events', [])
        if event_name in match_events:
            return vuln
        # 模糊匹配
        for me in match_events:
            if me in event_name or event_name in me:
                return vuln
    return None


def _assess_cvss_from_kb(event_name: str) -> Dict:
    """从知识库获取CVSS评分，未命中则关键词兜底"""
    vuln = _lookup_vuln(event_name)
    if vuln:
        cvss = vuln.get('cvss', 5.5)
        rec = vuln.get('recommendation_level', '建议')
        return {'cvss': cvss, 'risk_level': vuln.get('severity', '中'), 'recommendation': rec}

    # 关键词兜底
    high_kw = ['rce', '注入', '命令执行', '代码执行', '反序列化', 'sql注入', 'xss', '文件上传']
    event_lower = event_name.lower()
    for kw in high_kw:
        if kw in event_lower:
            return {'cvss': 8.5, 'risk_level': '高', 'recommendation': '强烈建议'}
    return {'cvss': 5.5, 'risk_level': '中', 'recommendation': '建议'}


# ============================================================
# 样式工具函数
# ============================================================

def set_chinese_font(document):
    """设置文档中文字体"""
    document.styles['Normal'].font.name = '宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_heading(document, text, level=1):
    """添加中文标题"""
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return heading


def add_paragraph(document, text):
    """添加中文段落"""
    para = document.add_paragraph(text)
    for run in para.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(11)
    return para


def add_bold_paragraph(document, text):
    """添加粗体段落"""
    para = document.add_paragraph()
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)
    run.bold = True
    return para


def add_table(document, rows, cols, data):
    """添加表格"""
    table = document.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'

    for i in range(rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = str(data[i][j])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)
    return table


def generate_summary_table(result: Dict) -> List[List]:
    """生成分析总览表格数据 — v2 基于 FP 统计"""
    summary = result['summary']
    fp_stats = result.get('weibu_fp_stats', {})
    fp_analysis = result.get('weibu_false_positive_analysis', {})

    total = (summary['lvmeng_unique_four_tuples'] +
             summary['weibu_unique_four_tuples'] +
             summary['common_four_tuples'])

    return [
        ['统计项', '数量/说明'],
        ['绿盟总告警数', str(summary['lvmeng_total_alerts'])],
        ['微步总告警数', str(summary['weibu_total_alerts'])],
        ['绿盟独有告警', str(summary['lvmeng_only_alert_count'])],
        ['微步独有告警', str(summary['weibu_only_alert_count'])],
        ['├─ 其中误报', f'{fp_stats.get("false_positive", 0)} ({fp_analysis.get("false_positive_rate", 0)}%)'],
        ['├─ 其中真阳性', str(fp_stats.get('true_positive', 0))],
        ['├─ 其中需人工审核', str(fp_stats.get('manual_review', 0))],
        ['交叉告警', str(summary['cross_alert_count'])],
        ['共同四元组数', str(summary['common_four_tuples'])],
    ]


def extract_ip(ip_str: str) -> str:
    """
    提取纯IP地址，去除括号和内部网段标注

    Args:
        ip_str: 原始IP字符串，如 "59.66.129.211(内部网段02)"

    Returns:
        纯IP地址，如 "59.66.129.211"
    """
    if not ip_str:
        return ''
    # 移除括号及其内容
    return re.sub(r'\(.*?\)', '', ip_str).strip()


def generate_report(result: Dict, output_path: str, unpaired_data: Dict = None):
    """
    生成Word报告

    Args:
        result: 对比分析结果
        output_path: 输出文件路径
        unpaired_data: 未配对厂商数据（厂商优势项）
    """
    document = Document()
    set_chinese_font(document)

    # 标题
    title = add_heading(document, '网络安全设备告警对比分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 一、分析总览
    add_heading(document, '一、分析总览', 1)

    summary_table_data = generate_summary_table(result)
    add_table(document, len(summary_table_data), 2, summary_table_data)

    # 二、绿盟独有告警分析
    if result['lvmeng_only_alerts']:
        add_heading(document, '二、绿盟独有告警分析', 1)

        # 按事件类型分组
        event_type_groups = {}
        for alert in result['lvmeng_only_alerts']:
            event_type = alert.get('event_type', '未知')
            if event_type not in event_type_groups:
                event_type_groups[event_type] = []
            event_type_groups[event_type].append(alert)

        sub_heading_counter = 1
        for event_type, alerts in event_type_groups.items():
            add_heading(document, f'2.{sub_heading_counter} {event_type}', 2)
            add_paragraph(document, f'绿盟检测到 {len(alerts)} 条该类型告警，微步未检测到。')

            # 按告警名称去重
            seen_event_names = set()
            unique_alerts = []
            for alert in alerts:
                event_name = alert.get('event_name', '未知')
                if event_name not in seen_event_names:
                    seen_event_names.add(event_name)
                    unique_alerts.append(alert)

            # 每个告警名称展开三级标题
            sub_sub_heading_counter = 1
            for alert in unique_alerts[:3]:  # 最多分析3个示例
                event_name = alert.get('event_name', '未知')
                add_heading(document, f'2.{sub_heading_counter}.{sub_sub_heading_counter} {event_name}', 3)

                # 添加四元组示例
                raw_four_tuple = alert.get('raw_four_tuple', ('', '', '', ''))
                src_ip = extract_ip(raw_four_tuple[0])
                src_port = raw_four_tuple[1]
                dst_ip = extract_ip(raw_four_tuple[2])
                dst_port = raw_four_tuple[3]
                example_str = f'{src_ip}:{src_port}->{dst_ip}:{dst_port}'
                generate_time = alert.get('generate_time', '')
                time_str = f'，时间：{generate_time}' if generate_time else ''
                add_paragraph(document, f'流量示例：{event_name}：{example_str}{time_str}')

                sub_sub_heading_counter += 1

            sub_heading_counter += 1

    # 三、微步独有告警详细分析与建议
    if result['weibu_only_alerts']:
        add_heading(document, '三、微步独有告警详细分析与建议', 1)

        # 按事件类型分组
        event_type_groups = {}
        for alert in result['weibu_only_alerts']:
            event_type = alert.get('event_type', '未知')
            if event_type not in event_type_groups:
                event_type_groups[event_type] = []
            event_type_groups[event_type].append(alert)

        sub_heading_counter = 1
        for event_type, alerts in event_type_groups.items():
            add_heading(document, f'3.{sub_heading_counter} {event_type}', 2)

            # 按告警名称去重，同一个告警名称只保留一条
            seen_event_names = set()
            unique_alerts = []
            for alert in alerts:
                event_name = alert.get('event_name', '未知')
                if event_name not in seen_event_names:
                    seen_event_names.add(event_name)
                    unique_alerts.append(alert)

            add_paragraph(document, f'微步检测到 {len(alerts)} 条该类型告警，绿盟未检测到。')

            # 分析每个告警（优先选择包含证据的告警作为示例）
            sub_sub_heading_counter = 1
            # 先筛选出有证据或编码绕过的告警
            priority_alerts = []
            other_alerts = []
            for alert in unique_alerts:
                fp = alert.get('false_positive_analysis', {})
                has_evidence = len(fp.get('evidence', [])) > 0
                has_bypass = fp.get('bypass_detection', {}).get('has_encoding', False)
                if has_evidence or has_bypass:
                    priority_alerts.append(alert)
                else:
                    other_alerts.append(alert)

            # 合并列表，优先显示有证据的
            display_alerts = priority_alerts + other_alerts

            for alert in display_alerts[:3]:  # 最多分析3个示例
                event_name = alert.get('event_name', '未知')
                cvss_info = _assess_cvss_from_kb(event_name)
                vuln = _lookup_vuln(event_name)

                add_heading(document,
                    f'3.{sub_heading_counter}.{sub_sub_heading_counter} {event_name}', 3)

                # 四元组示例
                raw_four_tuple = alert.get('raw_four_tuple', ('', '', '', ''))
                src_ip = extract_ip(raw_four_tuple[0])
                src_port = raw_four_tuple[1]
                dst_ip = extract_ip(raw_four_tuple[2])
                dst_port = raw_four_tuple[3]
                example_str = f'{src_ip}:{src_port}->{dst_ip}:{dst_port}'
                generate_time = alert.get('generate_time', '')
                time_str = f'，时间：{generate_time}' if generate_time else ''
                add_paragraph(document,
                    f'流量示例：{example_str}{time_str}')

                # 1. 漏洞基本信息
                add_bold_paragraph(document, '1. 漏洞基本信息')
                if vuln:
                    add_paragraph(document,
                        f'CVE编号: {vuln.get("cve", "N/A")}\n'
                        f'CVSS评分: {vuln.get("cvss", "N/A")}\n'
                        f'严重程度: {vuln.get("severity", "N/A")}\n'
                        f'影响系统: {vuln.get("affected", "N/A")}')
                else:
                    add_paragraph(document,
                        f'该告警属于 {event_type} 类型。')

                # 2. 攻击原理分析
                add_bold_paragraph(document, '2. 攻击原理分析')
                if vuln and vuln.get('principle'):
                    add_paragraph(document, vuln['principle'][:2000].strip())
                else:
                    add_paragraph(document,
                        f'该攻击类型属于{event_type}类攻击。'
                        '建议查阅相关CVE和威胁情报进行深入分析。')

                # 3. 绿盟检测能力缺口
                add_bold_paragraph(document, '3. 绿盟检测能力缺口')
                if vuln and vuln.get('detection_indicators'):
                    indicators = '\n'.join(
                        f'  - {ind}' for ind in vuln['detection_indicators'][:6])
                    add_paragraph(document,
                        f'绿盟当前未覆盖该漏洞的检测。微步能够检测到以下特征：\n{indicators}')
                else:
                    add_paragraph(document,
                        f'绿盟当前未覆盖该类型攻击的检测规则。')

                # 4. 误报分析
                if 'false_positive_analysis' in alert:
                    fp_result = alert['false_positive_analysis']
                    add_bold_paragraph(document, '4. 误报分析')
                    add_paragraph(document,
                        f'置信度：{fp_result.get("confidence", "未评估")}')

                    is_fp = fp_result.get('is_false_positive')
                    needs_review = fp_result.get('requires_manual_review')
                    if needs_review:
                        add_paragraph(document,
                            '是否误报：需人工审核')
                    elif is_fp is True:
                        add_paragraph(document,
                            '是否误报：是（正常业务流量）')
                    elif is_fp is False:
                        add_paragraph(document,
                            '是否误报：否（真阳性—确认安全威胁）')
                    else:
                        add_paragraph(document,
                            '是否误报：待确定')

                    if fp_result.get('false_positive_reason'):
                        add_paragraph(document,
                            f'判定理由：{fp_result["false_positive_reason"]}')

                    bypass = fp_result.get('bypass_detection', {})
                    if bypass.get('has_encoding'):
                        add_paragraph(document,
                            f'编码绕过检测：是 ({" + ".join(bypass.get("encoding_types", []))})')
                        if bypass.get('is_attack'):
                            add_paragraph(document,
                                '  解码后确认攻击：是')
                    else:
                        add_paragraph(document, '编码绕过检测：否')

                    if fp_result.get('evidence'):
                        add_paragraph(document, '判定证据：')
                        for ev in fp_result['evidence'][:8]:
                            add_paragraph(document, f'  - {ev}')

                # 5. 改进建议
                add_bold_paragraph(document, '5. 改进建议')
                rec_level = cvss_info['recommendation']
                add_paragraph(document,
                    f'建议等级：【{rec_level}】')
                if vuln:
                    add_paragraph(document,
                        f'理由：{vuln.get("recommendation_reason", "该漏洞具有安全风险")}')
                    if vuln.get('rule_example'):
                        add_paragraph(document, '检测规则示例：')
                        add_paragraph(document, vuln['rule_example'])
                else:
                    add_paragraph(document,
                        f'理由：{event_type}类型攻击具有一定安全风险，建议关注。'
                        f'建议针对{event_name}的攻击特征编写检测规则。')

                sub_sub_heading_counter += 1

                sub_sub_heading_counter += 1

            sub_heading_counter += 1

    # 四、交叉告警分析
    if result['cross_alerts']:
        add_heading(document, '四、交叉告警分析', 1)
        add_paragraph(document, f'双方均检测到 {len(result["cross_alerts"])} 个四元组的告警。')

        # 检测一致性的统计
        consistent_count = 0
        inconsistent_count = 0
        for comp in result['detailed_comparison']:
            if comp['consistency']['event_name_match']:
                consistent_count += 1
            else:
                inconsistent_count += 1

        add_paragraph(document, f'其中，告警判定一致的：{consistent_count} 个')
        add_paragraph(document, f'告警判定存在差异的：{inconsistent_count} 个')

        # 显示差异示例
        if inconsistent_count > 0:
            add_heading(document, '4.1 告警判定差异示例', 2)
            for comp in result['detailed_comparison'][:3]:
                if not comp['consistency']['event_name_match']:
                    four_tuple = comp['four_tuple']
                    lvmeng_names = list(set(a['event_name'] for a in comp['lvmeng_alerts']))
                    weibu_names = list(set(a['event_name'] for a in comp['weibu_alerts']))

                    add_paragraph(document, f'四元组：{four_tuple[0]}:{four_tuple[1]} -> {four_tuple[2]}:{four_tuple[3]}')
                    add_paragraph(document, f'  绿盟判定：{", ".join(lvmeng_names)}')
                    add_paragraph(document, f'  微步判定：{", ".join(weibu_names)}')
                    add_paragraph(document, '')

    # 五、总结与建议
    add_heading(document, '五、总结与建议', 1)

    summary = result['summary']
    if summary['weibu_only_alert_count'] > summary['lvmeng_only_alert_count']:
        add_paragraph(document, '总体分析：微步在本次检测中覆盖了更多类型的告警事件。')
        add_paragraph(document, '建议：')
        add_paragraph(document, '1. 关注微步独有的高、中风险告警类型，评估是否需要在绿盟中添加相应的检测规则。')
        add_paragraph(document, '2. 对于交叉告警中判定不一致的情况，建议人工复核并统一检测标准。')
    elif summary['lvmeng_only_alert_count'] > summary['weibu_only_alert_count']:
        add_paragraph(document, '总体分析：绿盟在本次检测中覆盖了更多类型的告警事件。')
    else:
        add_paragraph(document, '总体分析：两个厂商的检测覆盖范围基本相当。')

    # 六、厂商优势分析（未配对文件）
    if unpaired_data or result.get('unpaired_vendor_data'):
        vendor_advantage_data = unpaired_data or result.get('unpaired_vendor_data', {})
        add_heading(document, '六、厂商优势分析', 1)
        add_paragraph(document, '以下为未配对文件的分析结果，展示了各厂商在额外时间段的检测能力。')

        sub_heading_counter = 1
        for vendor_name, vendor_files in vendor_advantage_data.items():
            if not vendor_files:
                continue

            add_heading(document, f'6.{sub_heading_counter} {vendor_name}优势分析', 2)

            # 汇总统计
            total_alerts = 0
            all_event_types = {}

            for file_data in vendor_files:
                analysis = file_data.get('analysis', {})
                total_alerts += analysis.get('total_alerts', 0)

                for event_type, count in analysis.get('event_types', {}).items():
                    if event_type not in all_event_types:
                        all_event_types[event_type] = 0
                    all_event_types[event_type] += count

            # 统计摘要
            add_paragraph(document, '统计摘要：')
            add_paragraph(document, f'文件数量：{len(vendor_files)}')
            add_paragraph(document, f'总告警数：{total_alerts}')

            if all_event_types:
                add_paragraph(document, '主要事件类型：')
                sorted_types = sorted(all_event_types.items(), key=lambda x: x[1], reverse=True)
                for event_type, count in sorted_types[:10]:
                    add_paragraph(document, f'  • {event_type}：{count} 条')

            # 关键告警示例
            add_paragraph(document, '关键告警示例：')
            for file_data in vendor_files[:2]:  # 最多显示2个文件
                analysis = file_data.get('analysis', {})
                top_alerts = analysis.get('top_alerts', [])

                add_paragraph(document, f'\n文件：{file_data.get("filename", "未知")}')

                for alert in top_alerts[:5]:  # 每个文件最多显示5条
                    add_paragraph(document, f'1. {alert.get("event_name", "未知")} - {alert.get("threat_level", "未知")}')
                    add_paragraph(document, f'   事件类型：{alert.get("event_type", "未知")}')
                    add_paragraph(document, f'   四元组：{alert.get("source_ip", "")}:{alert.get("source_port", "")} -> {alert.get("destination_ip", "")}:{alert.get("destination_port", "")}')
                    add_paragraph(document, '')

            sub_heading_counter += 1

    # 保存文档
    document.save(output_path)
    print(f"报告已生成：{output_path}")


def main():
    if len(sys.argv) < 3:
        print("使用方法: python generate_report.py <分析结果JSON> <输出文件路径>")
        print("示例: python generate_report.py compare_result.json F:/claude/qinghua/report/分析报告.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_path = sys.argv[2]

    with open(input_file, 'r', encoding='utf-8') as f:
        result = json.load(f)

    generate_report(result, output_path)


if __name__ == "__main__":
    main()